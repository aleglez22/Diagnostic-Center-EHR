import re
import os
from docx import Document
from django.conf import settings

def docx_replace_regex(doc_obj, regex , replace):

    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex , replace)


def reemplaza(a_reemplazar,reemplazo,conclusiones,pacient_name,age,doct,date,doc_name ):
    c_campo = re.compile(r"{0}".format(str('xcampo')))
    campo = r"{0}".format(str(reemplazo))

    c_nombre_paciente = re.compile(r"{0}".format('xnombre'))
    nombre_paciente = r"{0}".format(str(pacient_name))

    c_conclusion = re.compile(r"{0}".format('xconclusion'))
    conclu = r"{0}".format(str(conclusiones))

    c_edad = re.compile(r"{0}".format('xedad'))
    edad = r"{0}".format(str(age))

    c_doctor= re.compile(r"{0}".format('xdoctor'))
    doctor = r"{0}".format(str(doct))

    c_fecha = re.compile(r"{0}".format('xfecha'))
    fecha = r"{0}".format(str(date))

    
    

    filename =str(doc_name)
    #module_dir=os.path.dirname(__file__)
    file_path =os.path.join((settings.MEDIA_ROOT), filename)
    doc = Document(file_path)

    docx_replace_regex(doc, c_campo, campo)
    docx_replace_regex(doc, c_nombre_paciente, nombre_paciente)
    docx_replace_regex(doc, c_edad, edad)
    docx_replace_regex(doc, c_doctor, doctor)
    docx_replace_regex(doc, c_fecha, fecha)
    docx_replace_regex(doc, c_conclusion, conclu)

    return doc





