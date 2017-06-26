
#para acceder a media
import os
from django.conf import settings
text = open(os.path.join(settings.MEDIA_ROOT, 'a.txt'), 'rb').read()

#para descargar el archivo de word

from django.http import HttpResponse
from docx import Document
from cStringIO import StringIO

def your_view(request):
    document = Document()
    document.add_heading(u"My title", 0)
    # add more things to your document with python-docx

    f = StringIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=example.docx'
    response['Content-Length'] = length
    return response


#mas limpia
from docx import Document
from django.http import HttpResponse

def download_docx(request):
    document = Document()
    document.add_heading('Document Title', 0)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=download.docx'
    document.save(response)

    return response