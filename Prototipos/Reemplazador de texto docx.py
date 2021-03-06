import re
import os
from docx import Document

def docx_replace_regex(doc_obj, regex , replace):

    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex , replace)


a_reemplazar='hola'
reemplazo='adios'
regex1 = re.compile(r"{0}".format(str(a_reemplazar)))
replace1 = r"{0}".format(str(reemplazo))
filename = "test.docx"
module_dir=os.path.dirname(__file__)
file_path =os.path.join(module_dir, filename)
doc = Document(file_path)
print(file_path)
docx_replace_regex(doc, regex1 , replace1)
doc.save('test.docx')

